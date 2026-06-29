from docx import Document as DocxDocument

from .models import Document, Page
import os


class DOCXParser:

    def parse(self, file_path):

        doc = DocxDocument(file_path)

        text = ""

        for para in doc.paragraphs:
            text += para.text + "\n"

        page = Page(
            page_number=1,
            text=text
        )

        return Document(
            file_name=os.path.basename(file_path),
            file_type="docx",
            pages=[page],
            metadata={}
        )